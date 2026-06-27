"""文档解析工具 — 提取 PDF/DOCX/XLSX/CSV/TXT 的文字和表格内容。"""

import csv
import logging
from pathlib import Path

logger = logging.getLogger("nano_agent.tools.document_parse")


class DocumentParser:
    TOOLS = [
        ("parse_document", "Parse and extract text/data from uploaded documents: "
         "PDF, Word (.docx), Excel (.xlsx), CSV, and plain text. "
         "Returns the full text content for the agent to analyze.",
         "parse_document",
         {"path": {"type": "string", "description": "Path to the document file in workspace"}},
         ["path"]),
        ("convert_to_pdf", "Convert a document or image to PDF. "
         "Supports Word (.docx), Excel (.xlsx), PowerPoint (.pptx), images (PNG/JPG/GIF/WEBP), "
         "and plain text files. Returns the path to the generated PDF.",
         "convert_to_pdf",
         {"path": {"type": "string", "description": "Path to the file to convert"},
          "output_name": {"type": "string", "description": "Optional output filename (default: auto-generated)"}},
         ["path"]),
    ]

    def __init__(self, work_dir: str):
        self.work_dir = Path(work_dir).resolve()

    def parse_document(self, path: str) -> str:
        """解析文档：自动检测类型，提取全部文字内容。"""
        try:
            filepath = (self.work_dir / path).resolve()
            filepath.relative_to(self.work_dir)
        except (ValueError, OSError):
            return "Error: Access denied"

        if not filepath.exists():
            return f"Error: File not found — {path}"

        ext = filepath.suffix.lower()
        size_mb = filepath.stat().st_size / (1024 * 1024)

        try:
            if ext == ".pdf":
                return self._parse_pdf(filepath)
            elif ext in (".docx", ".doc"):
                return self._parse_docx(filepath)
            elif ext in (".xlsx", ".xls"):
                return self._parse_excel(filepath)
            elif ext == ".csv":
                return self._parse_csv(filepath)
            elif ext in (".txt", ".md", ".py", ".json", ".xml", ".html", ".htm",
                         ".yaml", ".yml", ".toml", ".ini", ".cfg", ".log"):
                return self._parse_text(filepath)
            else:
                # 尝试当纯文本读
                try:
                    return self._parse_text(filepath)
                except Exception:
                    return f"Error: Unsupported format '{ext}'. Supported: PDF, DOCX, XLSX, CSV, TXT"
        except Exception as e:
            logger.error(f"Parse error for {path}: {e}")
            return f"Error parsing document: {e}"

    def _parse_pdf(self, filepath: Path) -> str:
        """解析 PDF，优先 pdfplumber，回退 PyPDF2。"""
        text_parts = []
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t:
                        text_parts.append(f"--- Page {i+1} ---\n{t}")
            if text_parts:
                return f"[PDF: {filepath.name}]\n\n" + "\n\n".join(text_parts)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # 回退 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            for i, page in enumerate(reader.pages):
                t = page.extract_text()
                if t:
                    text_parts.append(f"--- Page {i+1} ---\n{t}")
            return f"[PDF: {filepath.name}]\n\n" + "\n\n".join(text_parts)
        except Exception as e:
            return f"Error reading PDF: {e}"

    def _parse_docx(self, filepath: Path) -> str:
        """解析 Word 文档。"""
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # 也提取表格
            for ti, table in enumerate(doc.tables):
                text_parts.append(f"\n[Table {ti+1}]")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(cells))
            return f"[DOCX: {filepath.name}]\n\n" + "\n".join(text_parts)
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"
        except Exception as e:
            return f"Error reading DOCX: {e}"

    def _parse_excel(self, filepath: Path) -> str:
        """解析 Excel 表格，输出为 markdown 表格。"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, data_only=True, read_only=True)
            text_parts = [f"[Excel: {filepath.name}]"]

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 100:
                        text_parts.append(f"(Showing first 100 rows, {ws.max_row or '?'} total)")
                        break
                    if any(c is not None for c in row):
                        rows.append(row)

                if not rows:
                    continue

                text_parts.append(f"\n## Sheet: {sheet_name}")

                # 格式化为 markdown table
                max_cols = max(len(r) for r in rows)
                for i, row in enumerate(rows):
                    cells = [str(c) if c is not None else "" for c in row]
                    # Pad to max_cols
                    cells += [""] * (max_cols - len(cells))
                    text_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        text_parts.append("|" + "|".join(["---"] * max_cols) + "|")

            wb.close()
            return "\n".join(text_parts)
        except ImportError:
            return "Error: openpyxl not installed. Run: pip install openpyxl"
        except Exception as e:
            return f"Error reading Excel: {e}"

    def _parse_csv(self, filepath: Path) -> str:
        """解析 CSV 文件。"""
        try:
            with open(filepath, encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = list(reader)

            text_parts = [f"[CSV: {filepath.name}]\n"]
            if len(rows) > 200:
                text_parts.append(f"(Showing first 200 of {len(rows)} rows)")
                rows = rows[:200]

            max_cols = max(len(r) for r in rows) if rows else 0
            for i, row in enumerate(rows):
                cells = [str(c) for c in row]
                cells += [""] * (max_cols - len(cells))
                text_parts.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    text_parts.append("|" + "|".join(["---"] * max_cols) + "|")

            return "\n".join(text_parts)
        except UnicodeDecodeError:
            return "Error: CSV encoding not supported (try UTF-8)"
        except Exception as e:
            return f"Error reading CSV: {e}"

    def _parse_text(self, filepath: Path) -> str:
        """解析纯文本文件。"""
        try:
            content = filepath.read_text(encoding="utf-8")
            if len(content) > 50000:
                content = content[:50000] + "\n...(truncated)"
            return f"[{filepath.suffix.upper().lstrip('.')}: {filepath.name}]\n\n{content}"
        except UnicodeDecodeError:
            try:
                content = filepath.read_text(encoding="gbk")
                return f"[{filepath.name}]\n\n{content}"
            except Exception:
                return f"Error: Cannot decode file as text"

    def convert_to_pdf(self, path: str, output_name: str = "") -> str:
        """将文档或图片转为 PDF。支持 DOCX/XLSX/PPTX/图片/文本。"""
        try:
            filepath = (self.work_dir / path).resolve()
            filepath.relative_to(self.work_dir)
        except (ValueError, OSError):
            return "Error: Access denied"

        if not filepath.exists():
            return f"Error: File not found — {path}"

        ext = filepath.suffix.lower()
        out_name = output_name or filepath.stem
        out_name = out_name.replace(" ", "_").replace("/", "_")
        out_path = (self.work_dir / f"{out_name}.pdf").resolve()

        try:
            # Office 文档 → LibreOffice
            if ext in (".docx", ".doc", ".xlsx", ".xls", ".pptx", ".ppt", ".odt", ".ods"):
                return self._libreoffice_to_pdf(filepath, out_path)

            # 图片 → img2pdf
            elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"):
                return self._image_to_pdf(filepath, out_path)

            # 纯文本 → reportlab
            elif ext in (".txt", ".md", ".py", ".json", ".csv", ".log", ".yaml", ".yml"):
                return self._text_to_pdf(filepath, out_path)

            else:
                return f"Error: Unsupported format '{ext}' for PDF conversion. Supported: DOCX, XLSX, PPTX, PNG, JPG, TXT"
        except Exception as e:
            logger.error(f"PDF conversion error for {path}: {e}")
            return f"Error converting to PDF: {e}"

    def _libreoffice_to_pdf(self, filepath, out_path) -> str:
        """使用 LibreOffice headless 转换 Office 文档。"""
        import subprocess, shutil
        libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
        if not libreoffice:
            return "Error: LibreOffice not installed. Run: yum install libreoffice-headless"

        out_dir = str(self.work_dir)
        result = subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, str(filepath)],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            return f"Error: LibreOffice conversion failed: {result.stderr[:200]}"

        # LibreOffice 自动命名: xxx.pdf
        generated = self.work_dir / f"{filepath.stem}.pdf"
        if generated.exists():
            if str(generated) != str(out_path):
                generated.rename(out_path)
            return (
                f"✅ Converted to PDF: [{out_path.name}](/api/download/{out_path.name})\n"
                f"![PDF]({out_path.name})\n"
                f"[下载 {out_path.name}](/api/download/{out_path.name})"
            )
        return "Error: PDF file not generated by LibreOffice"

    def _image_to_pdf(self, filepath, out_path) -> str:
        """使用 img2pdf 将图片转为 PDF。"""
        try:
            import img2pdf
            with open(out_path, "wb") as f:
                f.write(img2pdf.convert(str(filepath)))
            return (
                f"✅ Converted to PDF: [{out_path.name}](/api/download/{out_path.name})\n"
                f"[下载 {out_path.name}](/api/download/{out_path.name})"
            )
        except ImportError:
            # 回退 PIL
            try:
                from PIL import Image
                img = Image.open(filepath)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(out_path, "PDF")
                return (
                    f"✅ Converted to PDF: [{out_path.name}](/api/download/{out_path.name})\n"
                    f"[下载 {out_path.name}](/api/download/{out_path.name})"
                )
            except ImportError:
                return "Error: img2pdf and PIL not installed. Run: pip install img2pdf"

    def _text_to_pdf(self, filepath, out_path) -> str:
        """使用 reportlab 将文本文件转为 PDF。"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import mm
        except ImportError:
            return "Error: reportlab not installed. Run: pip install reportlab"

        content = filepath.read_text(encoding="utf-8", errors="replace")
        doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                               leftMargin=20*mm, rightMargin=20*mm,
                               topMargin=15*mm, bottomMargin=15*mm)
        styles = getSampleStyleSheet()
        story = [Paragraph(f"<b>File: {filepath.name}</b>", styles["Heading2"])]

        for line in content.split("\n")[:500]:
            line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not line.strip():
                story.append(Paragraph("<br/>", styles["Normal"]))
            else:
                story.append(Paragraph(line, styles["Code"] if line.startswith((" ", "\t")) else styles["Normal"]))

        try:
            doc.build(story)
            return (
                f"✅ Converted to PDF: [{out_path.name}](/api/download/{out_path.name})\n"
                f"[下载 {out_path.name}](/api/download/{out_path.name})"
            )
        except Exception as e:
            return f"Error building PDF: {e}"
